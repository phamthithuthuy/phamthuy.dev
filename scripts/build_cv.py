#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Sinh CV chuẩn Harvard cho Phạm Thị Thu Thủy (Business Analyst), thân thiện ATS.

Usage:
  python3 scripts/build_cv.py
  # → public/cv_phamthuy.docx (+ .pdf nếu có LibreOffice)
"""
from pathlib import Path

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = Path(__file__).resolve().parent.parent
# DOCX outside public/ — Astro/Vite would try to parse binary assets in public/
OUT_DOCX = ROOT / "scripts" / "cv_phamthuy.docx"
OUT_PDF = ROOT / "public" / "cv_phamthuy.pdf"
OUT_PDF_DIR = OUT_PDF.parent

FONT = "Times New Roman"
USABLE_W = Cm(18.0)  # A4 21cm - lề trái 1.5 - lề phải 1.5

doc = Document()

# ---- Lề trang ----
sec = doc.sections[0]
sec.page_height = Cm(29.7)
sec.page_width = Cm(21.0)
sec.top_margin = Cm(1.2)
sec.bottom_margin = Cm(1.2)
sec.left_margin = Cm(1.5)
sec.right_margin = Cm(1.5)

# ---- Style mặc định ----
normal = doc.styles["Normal"]
normal.font.name = FONT
normal.font.size = Pt(10.5)
normal.element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
pf = normal.paragraph_format
pf.space_before = Pt(0)
pf.space_after = Pt(0)
pf.line_spacing = 1.0


def add_para(space_after=0, space_before=0, align=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(space_before)
    if align is not None:
        p.alignment = align
    return p


def run(p, text, *, bold=False, italic=False, size=10.5, color=None, caps=False):
    r = p.add_run(text)
    r.bold = bold
    r.italic = italic
    r.font.name = FONT
    r.font.size = Pt(size)
    r.element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    if color:
        r.font.color.rgb = RGBColor(*color)
    if caps:
        r.font.all_caps = True
    return r


def bottom_border(p, sz=8, color="000000"):
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(sz))
    bottom.set(qn("w:space"), "2")
    bottom.set(qn("w:color"), color)
    pBdr.append(bottom)
    pPr.append(pBdr)


def heading(text):
    p = add_para(space_before=5, space_after=3)
    run(p, text.upper(), bold=True, size=11)
    p.paragraph_format.tab_stops.clear_all()
    bottom_border(p)


def entry(left, right, *, left_bold=True, italic=False, size=10.5):
    """Một dòng: nội dung trái + (tab phải) ngày/thông tin bên phải."""
    p = add_para(space_before=3 if left_bold else 0, space_after=0)
    p.paragraph_format.tab_stops.add_tab_stop(USABLE_W, WD_TAB_ALIGNMENT.RIGHT)
    run(p, left, bold=left_bold, italic=italic, size=size)
    if right:
        p.add_run("\t")
        run(p, right, italic=True, size=size)
    return p


def bullet(segments, space_after=0):
    """segments: list các (text, {bold/italic}) hoặc str."""
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Cm(0.55)
    p.paragraph_format.first_line_indent = Cm(-0.3)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = 1.0
    for seg in segments:
        if isinstance(seg, str):
            run(p, seg)
        else:
            txt, opt = seg
            run(p, txt, bold=opt.get("bold", False), italic=opt.get("italic", False))
    return p


def skill_line(label, value):
    p = add_para(space_after=1)
    run(p, label + ": ", bold=True)
    run(p, value)
    return p


B = lambda t: (t, {"bold": True})
I = lambda t: (t, {"italic": True})

# ===================== HEADER =====================
p = add_para(space_after=1, align=WD_ALIGN_PARAGRAPH.CENTER)
run(p, "PHẠM THỊ THU THỦY", bold=True, size=18)

p = add_para(space_after=1, align=WD_ALIGN_PARAGRAPH.CENTER)
run(p, "Business Analyst", italic=True, size=10.5, color=(0x33, 0x33, 0x33))

p = add_para(space_after=2, align=WD_ALIGN_PARAGRAPH.CENTER)
run(
    p,
    "Hà Nội, Việt Nam  |  0355 021 275  |  thuy@phamthuy.dev  |  "
    "github.com/phamthithuthuy  |  phamthuy.dev",
    size=10,
)

# ===================== ĐỊNH HƯỚNG NGHỀ NGHIỆP =====================
heading("Định hướng nghề nghiệp")
p = add_para(space_after=2)
run(p, "Ngắn hạn: ", bold=True)
run(
    p,
    "Tìm kiếm vị trí Business Analyst Intern để vận dụng kỹ năng khai thác yêu cầu, viết "
    "user story/use case, xây dựng tài liệu đặc tả và mô hình hoá quy trình; đồng thời tích "
    "luỹ kinh nghiệm làm việc với stakeholder và đội phát triển. ",
)
run(p, "Dài hạn: ", bold=True)
run(
    p,
    "Trở thành Business Analyst có kiến thức nghiệp vụ và nền tảng kỹ thuật vững, có khả "
    "năng phân tích bài toán phức tạp, đề xuất giải pháp phù hợp và làm cầu nối hiệu quả "
    "giữa doanh nghiệp với đội ngũ phát triển sản phẩm.",
)

# ===================== HỌC VẤN =====================
heading("Học vấn")
entry("Đại học · Chuyên ngành Công nghệ Thông tin", "Năm cuối · Dự kiến tốt nghiệp 2027")
entry(
    "Kỹ sư Công nghệ Thông tin — Hà Nội, Việt Nam",
    "GPA: 3.6 / 4.0",
    left_bold=False,
    italic=True,
)
bullet(
    [
        B("Sinh viên Xuất sắc & Sinh viên Giỏi"),
        " trong 2 kỳ học gần nhất; nhận học bổng khuyến khích học tập.",
    ]
)
# ===================== DỰ ÁN =====================
heading("Dự án")

entry(
    "ShopFlow — Hệ thống Bán hàng & Quản lý kho (github.com/hoangtuan2k5/shopflow)",
    "05/2026 – 06/2026",
)
entry(
    "Business Analyst (nhóm 3 người · dự án học tập từ case study)",
    "Jira · Agile/Scrum · UML · Figma · SQL",
    left_bold=False,
    italic=True,
)
bullet(
    [
        "Phân tích case study bán hàng–kho; thực hiện ",
        B("competitor analysis"),
        " (đối chiếu "
        "tính năng các giải pháp POS/quản lý kho) để chốt ",
        B("3 nhóm actor"),
        " và ",
        B("phạm vi MVP"),
        " cho 3 sprint.",
    ]
)
bullet(
    [
        "Chuyển yêu cầu thành ",
        B("user story"),
        " kèm ",
        B("acceptance criteria"),
        "; quản lý ",
        B("product backlog trên Jira"),
        "; tham gia sprint planning, refinement và review.",
    ]
)
bullet(
    [
        "Viết ",
        B("SRS"),
        " cho các ",
        B("module"),
        " (danh mục sản phẩm, tồn kho…) và các ",
        B("luồng nghiệp vụ"),
        " (tạo đơn, mô phỏng thanh toán, nhập hàng, hoàn hàng, cảnh báo sắp hết hàng); "
        "mô hình hoá bằng activity / sequence / state cho vòng đời đơn hàng (tách trạng "
        "thái thanh toán và giao hàng).",
    ]
)
bullet(
    [
        "Phác thảo ",
        B("wireframe trên Figma"),
        " cho các màn hình chính để thống nhất flow với team "
        "trước khi phát triển.",
    ]
)
bullet(
    [
        "Phối hợp với ",
        B("FE và BE"),
        " để làm rõ requirement trong sprint; thực hiện ",
        B("kiểm thử / UAT theo AC"),
        ", ghi nhận lỗi và nghiệm thu trước khi chốt sprint; soạn user "
        "manual & checklist nghiệm thu cuối dự án.",
    ]
)
bullet(
    [
        "Ứng dụng ",
        B("AI"),
        " hỗ trợ tóm tắt đối thủ, draft tài liệu/story và gợi ý wireframe; "
        "dùng ",
        B("skill viết Jira issue qua Jira MCP"),
        " để soạn issue nhanh hơn; ",
        B("tự rà soát"),
        " trước khi chốt với team.",
    ],
    space_after=3,
)

entry("phamthuy.dev — Business Analysis Notebook & Personal Site", "06/2026 – nay")
entry("Tác giả", "Documentation · UML · BPMN", left_bold=False, italic=True)
bullet(
    [
        "Hệ thống hoá kiến thức BA dạng notebook: kỹ thuật elicitation, viết ",
        B("user story / use case / SRS"),
        ", mô hình hoá quy trình (BPMN/UML), và phân tích dữ liệu cơ bản.",
    ]
)
bullet(
    [
        "Xây dựng personal site giới thiệu portfolio các case study phân tích nghiệp vụ và "
        "tài liệu dự án."
    ],
    space_after=2,
)

# ===================== KỸ NĂNG =====================
heading("Kỹ năng")
skill_line(
    "Phân tích nghiệp vụ",
    "Elicitation · Competitor analysis · User Story · Use Case · SRS/BRD · "
    "Acceptance Criteria · UAT",
)
skill_line("Mô hình hoá", "UML (Use Case, Activity, Sequence, State) · BPMN · ERD · Wireframe")
skill_line("Công cụ", "Jira · Confluence · Figma · Draw.io · Postman · Microsoft Office")
skill_line(
    "AI / automation",
    "Jira MCP (soạn issue/story) · hỗ trợ research & draft tài liệu",
)
skill_line("Dữ liệu & Kỹ thuật", "SQL · PostgreSQL · REST API (đọc hiểu) · OpenAPI")
skill_line(
    "Phương pháp & Kỹ năng mềm",
    "Agile/Scrum · Quản lý backlog · Giao tiếp · Làm việc nhóm · Tư duy phản biện",
)

# ===================== HOẠT ĐỘNG & ĐỊNH HƯỚNG =====================
heading("Hoạt động & Định hướng")
bullet(
    [
        "Đang đào sâu nghiệp vụ phân tích (BABOK), kỹ thuật viết tài liệu và mô hình hoá quy trình."
    ]
)
bullet(
    ["Nâng cao tiếng Anh kỹ thuật để đọc tài liệu và giao tiếp trong môi trường phần mềm."]
)

OUT_DOCX.parent.mkdir(parents=True, exist_ok=True)
OUT_PDF_DIR.mkdir(parents=True, exist_ok=True)
doc.save(OUT_DOCX)
print(f"Saved {OUT_DOCX}")

# Optional: LibreOffice → PDF (site serves only the PDF from public/)
import shutil
import subprocess

soffice = shutil.which("libreoffice") or shutil.which("soffice")
if soffice:
    subprocess.run(
        [
            soffice,
            "--headless",
            "--convert-to",
            "pdf",
            "--outdir",
            str(OUT_PDF_DIR),
            str(OUT_DOCX),
        ],
        check=True,
    )
    print(f"Saved {OUT_PDF}")
else:
    print("LibreOffice not found — skipped PDF export (docx only).")
